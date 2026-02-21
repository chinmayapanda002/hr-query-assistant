"""
Document Ingestion Tool
Handles PDF, DOCX, and TXT ingestion into ChromaDB vector store.
Supports large unstructured company documents.
"""

import os
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime

import chromadb
from chromadb.config import Settings
# from langchain_community.vectorstores import Chroma
from langchain_chroma import Chroma
# from langchain_anthropic import ChatAnthropic
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.schema import Document
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
from dotenv import load_dotenv

load_dotenv()

CHROMA_PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "./data/chroma_db")
CHROMA_COLLECTION_NAME = os.getenv("CHROMA_COLLECTION_NAME", "hr_policies")
DOCUMENTS_DIR = os.getenv("DOCUMENTS_DIR", "./data/documents")


class DocumentIngestionTool:
    """
    Ingests HR policy documents into a ChromaDB vector store.
    Supports PDF, DOCX, TXT formats with intelligent chunking.
    """

    def __init__(self):
        self.embedding_model = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2"
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            length_function=len,
        )
        os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)
        self.vector_store = Chroma(
            collection_name=CHROMA_COLLECTION_NAME,
            embedding_function=self.embedding_model,
            persist_directory=CHROMA_PERSIST_DIR,
        )
        print(f"✅ ChromaDB initialized at {CHROMA_PERSIST_DIR}")

    def _detect_category(self, filename: str, content: str) -> str:
        """Auto-detect document category from filename and content."""
        filename_lower = filename.lower()
        content_lower = content[:2000].lower()

        categories = {
            "leave_policy": ["leave", "vacation", "pto", "sick", "maternity", "paternity", "holiday"],
            "reimbursement": ["reimbursement", "expense", "travel allowance", "claim", "reimburse"],
            "insurance": ["insurance", "health", "medical", "dental", "vision", "coverage", "premium"],
            "onboarding": ["onboarding", "new hire", "orientation", "joining", "induction"],
            "payroll": ["payroll", "salary", "compensation", "bonus", "pay", "ctc"],
            "performance": ["performance", "appraisal", "review", "kpi", "goals", "okr"],
            "code_of_conduct": ["code of conduct", "ethics", "compliance", "harassment", "discrimination"],
            "remote_work": ["remote", "work from home", "wfh", "hybrid", "telecommute"],
            "benefits": ["benefits", "perks", "welfare", "provident fund", "gratuity"],
            "it_policy": ["it policy", "data security", "password", "vpn", "device"],
        }

        for category, keywords in categories.items():
            if any(kw in filename_lower or kw in content_lower for kw in keywords):
                return category

        return "general_policy"

    def _extract_text_from_pdf(self, filepath: str) -> str:
        """Extract text from PDF files."""
        try:
            import PyPDF2
            text = ""
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"⚠️ PDF extraction error for {filepath}: {e}")
            return ""

    def _extract_text_from_docx(self, filepath: str) -> str:
        """Extract text from DOCX files."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join([cell.text for cell in row.cells]) + "\n"
            return text
        except Exception as e:
            print(f"⚠️ DOCX extraction error for {filepath}: {e}")
            return ""

    def _extract_text(self, filepath: str) -> str:
        """Extract text based on file type."""
        ext = Path(filepath).suffix.lower()
        if ext == ".pdf":
            return self._extract_text_from_pdf(filepath)
        elif ext in [".docx", ".doc"]:
            return self._extract_text_from_docx(filepath)
        elif ext in [".txt", ".md"]:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            print(f"⚠️ Unsupported file type: {ext}")
            return ""

    def ingest_document(self, filepath: str, metadata: Optional[Dict] = None) -> Dict:
        """
        Ingest a single document into the vector store.
        Returns ingestion statistics.
        """
        filepath = str(filepath)
        filename = Path(filepath).name

        # Extract text
        print(f"📄 Processing: {filename}")
        text = self._extract_text(filepath)

        if not text.strip():
            return {"status": "error", "message": "No text extracted", "filename": filename}

        # Clean text
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Auto-detect category
        category = self._detect_category(filename, text)

        # Create chunks
        chunks = self.text_splitter.split_text(text)

        # Build LangChain Documents
        doc_id = hashlib.md5(filepath.encode()).hexdigest()
        documents = []
        for i, chunk in enumerate(chunks):
            doc_meta = {
                "source": filename,
                "doc_id": doc_id,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "category": category,
                "ingestion_date": datetime.utcnow().isoformat(),
                **(metadata or {}),
            }
            documents.append(Document(page_content=chunk, metadata=doc_meta))

        # Add to vector store
        self.vector_store.add_documents(documents)
        print(f"✅ Ingested {len(chunks)} chunks from {filename} [category: {category}]")

        return {
            "status": "success",
            "filename": filename,
            "category": category,
            "chunks": len(chunks),
            "char_count": len(text),
        }

    def ingest_directory(self, directory: str = DOCUMENTS_DIR) -> List[Dict]:
        """Ingest all documents from a directory."""
        results = []
        supported = [".pdf", ".docx", ".doc", ".txt", ".md"]
        dir_path = Path(directory)

        if not dir_path.exists():
            dir_path.mkdir(parents=True, exist_ok=True)
            print(f"📁 Created documents directory: {directory}")
            return results

        files = [f for f in dir_path.iterdir() if f.suffix.lower() in supported]
        print(f"📂 Found {len(files)} documents to ingest from {directory}")

        for filepath in files:
            result = self.ingest_document(str(filepath))
            results.append(result)

        return results

    def similarity_search(
        self,
        query: str,
        k: int = 5,
        category_filter: Optional[str] = None
    ) -> List[Document]:
        """Search for relevant documents."""
        if category_filter:
            results = self.vector_store.similarity_search(
                query,
                k=k,
                filter={"category": category_filter}
            )
        else:
            results = self.vector_store.similarity_search(query, k=k)
        return results

    def similarity_search_with_score(
        self,
        query: str,
        k: int = 5
    ) -> List[Tuple[Document, float]]:
        """Search with relevance scores."""
        return self.vector_store.similarity_search_with_score(query, k=k)

    def get_collection_stats(self) -> Dict:
        """Get vector store statistics."""
        collection = self.vector_store._collection
        count = collection.count()
        return {
            "total_chunks": count,
            "collection_name": CHROMA_COLLECTION_NAME,
            "persist_dir": CHROMA_PERSIST_DIR,
        }


# Singleton instance
_ingestion_tool: Optional[DocumentIngestionTool] = None


def get_ingestion_tool() -> DocumentIngestionTool:
    global _ingestion_tool
    if _ingestion_tool is None:
        _ingestion_tool = DocumentIngestionTool()
    return _ingestion_tool
